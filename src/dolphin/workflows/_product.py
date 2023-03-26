"""Module for creating the OPERA output product in NetCDF format."""
from io import StringIO
from itertools import groupby
from typing import Any, Dict, List, Optional, Sequence, Tuple

import h5netcdf
import h5py
import numpy as np
import pyproj
from numpy.typing import ArrayLike, DTypeLike

from dolphin import io
from dolphin._log import get_log
from dolphin._types import Filename

from ._pge_runconfig import RunConfig

logger = get_log(__name__)


BASE_GROUP_NAME = "/science/SENTINEL1"
DISP_GROUP_NAME = f"{BASE_GROUP_NAME}/DISP"
CORRECTIONS_GROUP_NAME = f"{DISP_GROUP_NAME}/corrections"
IDENTIFICATION_GROUP_NAME = f"{DISP_GROUP_NAME}/identification"
GLOBAL_ATTRS = dict(
    Conventions="CF-1.8",
    contact="operaops@jpl.nasa.gov",
    institution="NASA JPL",
    mission_name="OPERA",
    reference_document="TBD",
    title="OPERA L3_DISP_S1 Product",
)

# Convert chunks to a tuple or h5py errors
HDF5_OPTS = io.DEFAULT_HDF5_OPTIONS.copy()
HDF5_OPTS["chunks"] = tuple(HDF5_OPTS["chunks"])  # type: ignore
# The GRID_MAPPING_DSET variable is used to store the name of the dataset containing
# the grid mapping information, which includes the coordinate reference system (CRS)
# and the GeoTransform. This is in accordance with the CF 1.8 conventions for adding
# geospatial metadata to NetCDF files.
# http://cfconventions.org/cf-conventions/cf-conventions.html#grid-mappings-and-projections
# Note that the name "spatial_ref" used here is arbitrary, but it follows the default
# used by other libraries, such as rioxarray:
# https://github.com/corteva/rioxarray/blob/5783693895b4b055909c5758a72a5d40a365ef11/rioxarray/rioxarray.py#L34 # noqa
GRID_MAPPING_DSET = "spatial_ref"


def create_output_product(
    unw_filename: Filename,
    conncomp_filename: Filename,
    tcorr_filename: Filename,
    output_name: Filename,
    corrections: Dict[str, ArrayLike] = {},
    pge_runconfig: Optional[RunConfig] = None,
):
    """Create the OPERA output product in NetCDF format.

    Parameters
    ----------
    unw_filename : Filename
        The path to the input unwrapped phase image.
    conncomp_filename : Filename
        The path to the input connected components image.
    tcorr_filename : Filename
        The path to the input temporal correlation image.
    output_name : Filename, optional
        The path to the output NetCDF file, by default "output.nc"
    corrections : Dict[str, ArrayLike], optional
        A dictionary of corrections to write to the output file, by default None
    pge_runconfig : Optional[RunConfig], optional
        The PGE run configuration, by default None
        Used to add extra metadata to the output file.
    """
    # Read the Geotiff file and its metadata
    crs = io.get_raster_crs(unw_filename)
    gt = io.get_raster_gt(unw_filename)
    unw_arr = io.load_gdal(unw_filename)

    conncomp_arr = io.load_gdal(conncomp_filename)
    tcorr_arr = io.load_gdal(tcorr_filename)

    # Get the nodata mask (which for snaphu is 0)
    mask = unw_arr == 0
    # Set to NaN for final output
    unw_arr[mask] = np.nan

    assert unw_arr.shape == conncomp_arr.shape == tcorr_arr.shape

    with h5netcdf.File(output_name, "w") as f:
        # Create the NetCDF file
        f.attrs.update(GLOBAL_ATTRS)

        displacement_group = f.create_group(DISP_GROUP_NAME)

        # Set up the grid mapping variable for each group with rasters
        _create_grid_mapping(group=displacement_group, crs=crs, gt=gt)

        # Set up the X/Y variables for each group
        _create_yx_dsets(group=displacement_group, gt=gt, shape=unw_arr.shape)

        # Write the displacement array / conncomp arrays
        _create_geo_dataset(
            group=displacement_group,
            name="unwrapped_phase",
            data=unw_arr,
            description="Unwrapped phase",
            fillvalue=np.nan,
            attrs=dict(units="radians"),
        )
        _create_geo_dataset(
            group=displacement_group,
            name="connected_components",
            data=conncomp_arr,
            description="Connected components of the unwrapped phase",
            fillvalue=0,
            attrs=dict(units="unitless"),
        )
        _create_geo_dataset(
            group=displacement_group,
            name="temporal_correlation",
            data=tcorr_arr,
            description="Temporal correlation of phase inversion",
            fillvalue=np.nan,
            attrs=dict(units="unitless"),
        )

        # Create the group holding phase corrections that were used on the unwrapped phase
        corrections_group = f.create_group(CORRECTIONS_GROUP_NAME)
        f.attrs["description"] = "Phase corrections applied to the unwrapped_phase"

        # TODO: Are we going to downsample these for space?
        # if so, they need they're own X/Y variables and GeoTransform
        _create_grid_mapping(group=corrections_group, crs=crs, gt=gt)
        _create_yx_dsets(group=corrections_group, gt=gt, shape=unw_arr.shape)
        troposphere = corrections.get("troposphere", np.zeros_like(unw_arr))
        _create_geo_dataset(
            group=corrections_group,
            name="tropospheric_delay",
            data=troposphere,
            description="Tropospheric phase delay used to correct the unwrapped phase",
            fillvalue=np.nan,
            attrs=dict(units="radians"),
        )
        ionosphere = corrections.get("ionosphere", np.zeros_like(unw_arr))
        _create_geo_dataset(
            group=corrections_group,
            name="ionospheric_delay",
            data=ionosphere,
            description="Ionospheric phase delay used to correct the unwrapped phase",
            fillvalue=np.nan,
            attrs=dict(units="radians"),
        )
        solid_earth = corrections.get("solid_earth", np.zeros_like(unw_arr))
        _create_geo_dataset(
            group=corrections_group,
            name="solid_earth_tide",
            data=solid_earth,
            description="Solid Earth tide used to correct the unwrapped phase",
            fillvalue=np.nan,
            attrs=dict(units="radians"),
        )
        plate_motion = corrections.get("plate_motion", np.zeros_like(unw_arr))
        _create_geo_dataset(
            group=corrections_group,
            name="plate_motion",
            data=plate_motion,
            description="Phase ramp caused by plate",
            fillvalue=np.nan,
            attrs=dict(units="radians"),
        )
        # Make a scalar dataset for the reference point
        reference_point = corrections.get("reference_point", 0.0)
        _create_dataset(
            group=corrections_group,
            name="reference_point",
            dimensions=(),
            data=reference_point,
            fillvalue=np.nan,
            description=(
                "The constant phase subtracted from the unwrapped phase to"
                " zero-reference."
            ),
            dtype=np.float32,
            # Note: the dataset is only a scalar, but it could have come from multiple
            # points (e.g. some boxcar average of an area).
            # So the attributes will be lists of those locations used
            attrs=dict(units="radians", rows=[], cols=[], latitudes=[], longitudes=[]),
        )

    # End of the product for non-PGE users
    if pge_runconfig is None:
        return

    # Add the PGE metadata to the file
    with h5netcdf.File(output_name, "a") as f:
        identification_group = f.create_group(IDENTIFICATION_GROUP_NAME)
        _create_dataset(
            group=identification_group,
            name="frame_id",
            dimensions=(),
            data=pge_runconfig.input_file_group.frame_id,
            fillvalue=None,
            description="ID number of the processed frame.",
            attrs=dict(units="unitless"),
        )
        # TODO: prob should just make a _to_string method?
        ss = StringIO()
        pge_runconfig.to_yaml(ss)
        runconfig_str = ss.getvalue()
        _create_dataset(
            group=identification_group,
            name="pge_runconfig",
            dimensions=(),
            data=runconfig_str,
            fillvalue=None,
            description="The full PGE runconfig YAML file.",
            attrs=dict(units="unitless"),
        )


def _create_dataset(
    *,
    group: h5netcdf.Group,
    name: str,
    dimensions: Optional[Sequence[str]],
    data: np.ndarray,
    description: str,
    fillvalue: Optional[float],
    attrs: Optional[Dict[str, Any]] = None,
    dtype: Optional[DTypeLike] = None,
) -> h5netcdf.Variable:
    if attrs is None:
        attrs = {}
    attrs.update(long_name=description)

    # Scalars don't need chunks/compression
    if not isinstance(data, str) and np.array(data).size > 1:
        options = HDF5_OPTS
    else:
        options = {}
    dset = group.create_variable(
        name,
        dimensions=dimensions,
        data=data,
        dtype=dtype,
        fillvalue=fillvalue,
        **options,
    )
    dset.attrs.update(attrs)
    return dset


def _create_geo_dataset(
    *,
    group: h5netcdf.Group,
    name: str,
    data: np.ndarray,
    description: str,
    fillvalue: float,
    attrs: Optional[Dict[str, Any]],
) -> h5netcdf.Variable:
    dimensions = ["y", "x"]
    dset = _create_dataset(
        group=group,
        name=name,
        dimensions=dimensions,
        data=data,
        description=description,
        fillvalue=fillvalue,
        attrs=attrs,
    )
    dset.attrs["grid_mapping"] = GRID_MAPPING_DSET
    return dset


def _create_yx_arrays(
    gt: List[float], shape: Tuple[int, int]
) -> Tuple[np.ndarray, np.ndarray]:
    """Create the x and y coordinate datasets."""
    ysize, xsize = shape
    # Parse the geotransform
    x_origin, x_res, _, y_origin, _, y_res = gt

    # Make the x/y arrays
    # Note that these are the center of the pixels, whereas the GeoTransform
    # is the upper left corner of the top left pixel.
    y = np.arange(y_origin + y_res / 2, y_origin + y_res * ysize, y_res)
    x = np.arange(x_origin + x_res / 2, x_origin + x_res * xsize, x_res)
    return y, x


def _create_yx_dsets(
    group: h5netcdf.Group,
    gt: List[float],
    shape: Tuple[int, int],
) -> Tuple[h5netcdf.Variable, h5netcdf.Variable]:
    """Create the x and y coordinate datasets."""
    y, x = _create_yx_arrays(gt, shape)

    if not group.dimensions:
        group.dimensions = dict(y=y.size, x=x.size)
    # Create the datasets
    y_ds = group.create_variable("y", ("y",), data=y, dtype=float)
    x_ds = group.create_variable("x", ("x",), data=x, dtype=float)

    for name, ds in zip(["y", "x"], [y_ds, x_ds]):
        ds.attrs["standard_name"] = f"projection_{name}_coordinate"
        ds.attrs["long_name"] = f"{name.replace('_', ' ')} coordinate of projection"
        ds.attrs["units"] = "m"

    return y_ds, x_ds


def _create_grid_mapping(group, crs: pyproj.CRS, gt: List[float]) -> h5netcdf.Variable:
    """Set up the grid mapping variable."""
    # https://github.com/corteva/rioxarray/blob/21284f67db536d9c104aa872ab0bbc261259e59e/rioxarray/rioxarray.py#L34
    dset = group.create_variable(GRID_MAPPING_DSET, (), data=0, dtype=int)

    dset.attrs.update(crs.to_cf())
    # Also add the GeoTransform
    gt_string = " ".join([str(x) for x in gt])
    dset.attrs.update(
        dict(
            GeoTransform=gt_string,
            units="unitless",
            long_name=(
                "Dummy variable containing geo-referencing metadata in attributes"
            ),
        )
    )

    return dset


def _generate_docx_table(hdf5_path: Filename, output_path: Filename):
    # https://python-docx.readthedocs.io/en/latest/user/quickstart.html#adding-a-table
    from docx import Document
    from docx.enum.table import WD_ROW_HEIGHT_RULE, WD_TABLE_ALIGNMENT
    from docx.oxml import parse_xml
    from docx.oxml.ns import nsdecls
    from docx.shared import Pt

    def _add_row(table, text, height=15, shade=False, bold=False):
        # _tc.get_or_add_tcPr().append(shading_elm)
        row = table.add_row()
        row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
        row.height = Pt(height)
        row.cells[0].text = text
        # https://stackoverflow.com/questions/26752856/python-docx-set-table-cell-background-and-text-color  # noqa
        if shade:
            shading_elm = parse_xml(r'<w:shd {} w:fill="D9D9D9"/>'.format(nsdecls("w")))
            row.cells[0]._tc.get_or_add_tcPr().append(shading_elm)
        # Set the text color to black and remove bold
        run = row.cells[0].paragraphs[0].runs[0]
        run.font.color.rgb = None
        if not bold:
            run.font.bold = False

    document = Document()
    # Set the default document font to Arial
    style = document.styles["Normal"]
    font = style.font
    font.name = "Arial"

    for group_name, rows in _get_hdf5_attributes_by_group(hdf5_path).items():
        document.add_heading(f"Group: {group_name}", level=2)
        table = document.add_table(cols=1, rows=0)
        table.style = "Table Grid"  # Use the "Table Grid" style to get borders
        table.alignment = WD_TABLE_ALIGNMENT.LEFT

        for row in rows:
            name = row.pop("Name")
            desc = row.pop("Description")

            _add_row(table, f"Name: {name}", shade=True)

            row_text = "\t\t".join(f"{k}: {v or 'scalar'}" for k, v in row.items())
            row_text.replace("()", "scalar")
            _add_row(table, row_text)
            _add_row(table, f"Description: {desc}")

    logger.info(f"Saving to {output_path}")
    document.save(output_path)


def _get_hdf5_attributes(hdf5_path: Filename) -> List:
    table_data = []

    def append_dset_to_table(name, item):
        """Add all dataset's metadata using `visititems`."""
        if not isinstance(item, h5py.Dataset):
            return None
        data_type = item.dtype
        shape = item.shape
        description = item.attrs.get("long_name", "")
        units = item.attrs.get("units", "")
        table_data.append(
            dict(
                Name=name,
                Type=data_type,
                Shape=shape,
                Units=units,
                Description=description,
            )
        )

    with h5py.File(hdf5_path, "r") as hf:
        hf.visititems(append_dset_to_table)
    return table_data


def _get_hdf5_attributes_by_group(hdf5_path: Filename) -> Dict[str, List]:
    def get_group(name):
        return name.split("/")[-2]

    table_data = _get_hdf5_attributes(hdf5_path)

    group_sorted_rows = sorted(table_data, key=lambda row: get_group(row["Name"]))
    # Make a dict, where keys are group name, value is the list of rows
    # e.g.:  { 'DISP': [ {'Name': ,....], 'corrections': [{'Name':...}]
    return {
        k: list(v)
        for k, v in groupby(group_sorted_rows, key=lambda row: get_group(row["Name"]))
    }
