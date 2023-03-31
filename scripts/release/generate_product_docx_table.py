#!/usr/bin/env python
import argparse
from itertools import groupby
from typing import Dict, List

import h5py

from dolphin._types import Filename


def generate_docx_table(hdf5_path: Filename, output_path: Filename):
    """Create a Word document with a table of HDF5 datasets."""
    # https://python-docx.readthedocs.io/en/latest/user/quickstart.html#adding-a-table
    from docx import Document
    from docx.enum.table import WD_ROW_HEIGHT_RULE, WD_TABLE_ALIGNMENT
    from docx.oxml import parse_xml
    from docx.oxml.ns import nsdecls
    from docx.shared import Pt

    def _add_row(table, text, height=15, shade=False, bold=False):
        # _tc.get_or_add_tcPr().append(shading_elm)
        row = table.add_row()
        row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
        row.height = Pt(height)
        if isinstance(text, list):
            for i in range(len(text)):
                row.cells[i].text = text[i]
        else:
            row.cells[1].merge(row.cells[0])
            row.cells[1].merge(row.cells[2])
            row.cells[1].text = text
        # https://stackoverflow.com/questions/26752856/python-docx-set-table-cell-background-and-text-color  # noqa
        if shade:
            shading_elm = parse_xml(r'<w:shd {} w:fill="D9D9D9"/>'.format(nsdecls("w")))
            row.cells[0]._tc.get_or_add_tcPr().append(shading_elm)
        # Set the text color to black and remove bold
        run = row.cells[0].paragraphs[0].runs[0]
        run.font.color.rgb = None
        if not bold:
            run.font.bold = False

    document = Document()
    # Set the default document font to Arial
    style = document.styles["Normal"]
    font = style.font
    font.name = "Arial"

    for group_name, rows in _get_hdf5_attributes_by_group(hdf5_path).items():
        document.add_heading(f"Group: {group_name}", level=2)
        table = document.add_table(cols=3, rows=0)
        table.style = "Table Grid"  # Use the "Table Grid" style to get borders
        table.alignment = WD_TABLE_ALIGNMENT.LEFT

        for row in rows:
            name = row.pop("Name")
            desc = row.pop("Description")

            _add_row(table, f"Name: {name}", shade=True)

            row_text = [f"{k}: {v or 'scalar'}" for k, v in row.items()]
            _add_row(table, row_text)
            _add_row(table, f"Description: {desc}")

    print(f"Saving to {output_path}")
    document.save(output_path)


def _get_hdf5_attributes(hdf5_path: Filename) -> List:
    table_data = []

    def append_dset_to_table(name, item):
        """Add all dataset's metadata using `visititems`."""
        if not isinstance(item, h5py.Dataset):
            return None
        data_type = item.dtype
        shape = item.shape
        description = item.attrs.get("long_name", "")
        units = item.attrs.get("units", "")
        table_data.append(
            dict(
                Name=name,
                Type=data_type,
                Shape=shape,
                Units=units,
                Description=description,
            )
        )

    with h5py.File(hdf5_path, "r") as hf:
        hf.visititems(append_dset_to_table)
    return table_data


def _get_hdf5_attributes_by_group(hdf5_path: Filename) -> Dict[str, List]:
    def get_group(name):
        try:
            return name.split("/")[-2]
        except IndexError:
            return "root"

    table_data = _get_hdf5_attributes(hdf5_path)

    group_sorted_rows = sorted(table_data, key=lambda row: get_group(row["Name"]))
    # Make a dict, where keys are group name, value is the list of rows
    # e.g.:  { 'DISP': [ {'Name': ,....], 'corrections': [{'Name':...}]
    return {
        k: list(v)
        for k, v in groupby(group_sorted_rows, key=lambda row: get_group(row["Name"]))
    }


if __name__ == "__main__":
    parser = argparse.ArgumentParser()
    parser.add_argument("hdf5_path", type=str, help="Path to the HDF5 file")
    parser.add_argument(
        "output_path", type=str, help="Path to the output Word docx file"
    )
    args = parser.parse_args()
    generate_docx_table(args.hdf5_path, args.output_path)
